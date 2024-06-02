# El archivo frosty_app.py utiliza las funciones y constantes definidas en prompts.py para crear la 
# aplicaci칩n web interactiva de ExportBot. Importa la funci칩n get_system_prompt() para obtener el 
# mensaje del sistema inicial y lo utiliza para inicializar el historial de mensajes.
#####################################################################################################

# Se importan las bibliotecas necesarias (openai, re, streamlit) y la funci칩n get_system_prompt del 
# archivo prompts.py.

from openai import OpenAI  # Importa la biblioteca OpenAI para interactuar con modelos de lenguaje de OpenAI.
import re  # Importa el m칩dulo re para trabajar con expresiones regulares.
import streamlit as st  # Importa la biblioteca Streamlit para crear aplicaciones web interactivas.
from prompts import get_system_prompt  # Importa una funci칩n personalizada para obtener el mensaje del sistema.
import io
import pandas as pd

########################################
from docx import Document
from docx.shared import Inches
from io import BytesIO
from funciones import estilos, formato_tablas, formatear_numero
from docx.shared import Pt

# Definir nombres de estilo como constantes
TITLE_STYLE_NAME = 'MyTitleStyle'
HEADING1_STYLE_NAME = 'MyHeading1'
HEADING2_STYLE_NAME = 'MyHeading2'
TABLE_STYLE_NAME = 'MyTableStyle'
NORMAL_STYLE_NAME = 'Normal'

from PIL import Image
########################################

st.title("游뱄 ExportBot 游깵")  # Establece el t칤tulo de la aplicaci칩n web.

# Carga las im치genes de los logos
logo_mincit = Image.open("LogoMinCIT.png")
logo_procolombia = Image.open("LogoProColombia.png")

# Crea dos columnas para los logos
col1, col2 = st.columns(2)

# Muestra los logos en las columnas correspondientes
with col1:
    st.image(logo_mincit, width=200)
with col2:
    st.image(logo_procolombia, width=200)

# Se inicializa el historial de mensajes del chat. Se crea un cliente de OpenAI utilizando la clave API 
# almacenada en los secretos de Streamlit. Si "messages" no est치 en el estado de la sesi칩n, se inicializa 
# con el mensaje del sistema obtenido de la funci칩n get_system_prompt().

########################################
from funciones import estilos, formato_tablas, formatear_numero

from docx.enum.text import WD_ALIGN_PARAGRAPH

def to_word(df, prompt):
    """
    Genera un documento de Word con los resultados del chat.

    Args:
        df (pandas.DataFrame): DataFrame con los resultados del chat.
        prompt (str): Prompt utilizado para la consulta.

    Returns:
        BytesIO: Buffer de memoria que contiene el documento de Word generado.
    """
    # Crea un nuevo documento de Word
    doc = Document()

    # Agrega una tabla de 1 fila y 2 columnas para los logos
    logo_table = doc.add_table(rows=1, cols=2)
    logo_table.autofit = False
    logo_table.allow_autofit = False
    logo_table.columns[0].width = Inches(3)
    logo_table.columns[1].width = Inches(3)
    
    # Inserta el logo de MinCIT en la primera celda
    logo_mincit_path = "LogoMinCIT.png"
    logo_mincit_cell = logo_table.rows[0].cells[0]
    logo_mincit_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_mincit_run = logo_mincit_cell.paragraphs[0].add_run()
    logo_mincit_run.add_picture(logo_mincit_path, width=Inches(1.5))
    
    # Inserta el logo de ProColombia en la segunda celda
    logo_procolombia_path = "LogoProColombia.png"
    logo_procolombia_cell = logo_table.rows[0].cells[1]
    logo_procolombia_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_procolombia_run = logo_procolombia_cell.paragraphs[0].add_run()
    logo_procolombia_run.add_picture(logo_procolombia_path, width=Inches(1.5))
    
    # Agrega un salto de p치rrafo despu칠s de la tabla de logos
    doc.add_paragraph()

    # Aplica estilos al documento
    estilos(doc)

    # Agrega un t칤tulo al documento
    titulo = doc.add_heading('Resultados del Chat', level=0)
    titulo.style = TITLE_STYLE_NAME

    # Agrega el prompt como un p치rrafo
    prompt_text = f"Prompt: {prompt}"
    prompt_paragraph = doc.add_paragraph(prompt_text)
    prompt_paragraph.style = NORMAL_STYLE_NAME

    # Agrega una tabla al documento
    table = doc.add_table(rows=df.shape[0]+1, cols=df.shape[1])

    # Agrega los encabezados de la tabla
    for j in range(df.shape[1]):
        cell = table.cell(0, j)
        cell.text = str(df.columns[j])
        cell.paragraphs[0].runs[0].bold = True

    # Agrega los datos de la tabla
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = table.cell(i+1, j)
            cell.text = formatear_numero(df.iloc[i, j])

    # Formatea la tabla
    formato_tablas(doc, table)

    # Ajusta el espaciado antes y despu칠s de la tabla
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(12)
    table.rows[-1].cells[0].paragraphs[0].paragraph_format.space_after = Pt(12)

    # Agrega la advertencia como un p치rrafo al final del documento
    advertencia = """
    
    Advertencia: Revisar los resultados cuidadosamente. La informaci칩n generada por IA puede contener errores.
    
    La informaci칩n contenida en este documento es de orientaci칩n y gu칤a general. En ning칰n caso, ProColombia, ni sus empleados, son responsables ante usted o cualquier otra persona por las decisiones o acciones que pueda tomar en relaci칩n con la informaci칩n proporcionada, por lo cual debe tomarse como de car치cter referencial 칰nicamente.
    """
    advertencia_paragraph = doc.add_paragraph(advertencia)
    advertencia_paragraph.style = NORMAL_STYLE_NAME

    # Guarda el documento en un buffer de memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer
########################################

# Inicializa el historial de mensajes del chat.
client = OpenAI(api_key=st.secrets.OPENAI_API_KEY)  # Crea un cliente de OpenAI utilizando la clave API almacenada en los secretos de Streamlit.
if "messages" not in st.session_state:  # Verifica si "messages" no est치 en el estado de la sesi칩n.
    # El mensaje del sistema incluye informaci칩n de la tabla, reglas y genera un mensaje de bienvenida para el usuario.
    st.session_state.messages = [{"role": "system", "content": get_system_prompt()}]  # Inicializa "messages" con el mensaje del sistema.

# Se solicita la entrada del usuario y se guarda. Si el usuario ingresa un mensaje en el campo de entrada 
# del chat, se agrega al historial de mensajes.

# Solicita la entrada del usuario y la guarda.
if prompt := st.chat_input():  # Si el usuario ingresa un mensaje en el campo de entrada del chat.
    st.session_state.messages.append({"role": "user", "content": prompt})  # Agrega el mensaje del usuario al historial de mensajes.

# Se muestran los mensajes del chat existentes. Se itera sobre cada mensaje en el historial de mensajes, 
# omitiendo el mensaje del sistema. Se crea un contenedor para cada mensaje seg칰n el rol (usuario o 
# asistente) y se escribe el contenido del mensaje en la aplicaci칩n web. Si el mensaje contiene resultados
# de una consulta SQL, se muestran en un DataFrame.

# Muestra los mensajes del chat existentes.
for message in st.session_state.messages:  # Itera sobre cada mensaje en el historial de mensajes.
    if message["role"] == "system":  # Omite el mensaje del sistema.
        continue
    with st.chat_message(message["role"]):  # Crea un contenedor para el mensaje seg칰n el rol (usuario o asistente).
        st.write(message["content"])  # Escribe el contenido del mensaje en la aplicaci칩n web.
        if "results" in message:  # Si el mensaje contiene resultados de una consulta SQL.
            st.dataframe(message["results"])  # Muestra los resultados en un DataFrame.
    
####################################

####################################

# Si el 칰ltimo mensaje no es del asistente, se genera una nueva respuesta. Se crea un contenedor vac칤o para
# actualizar din치micamente la respuesta. Se llama a la API de OpenAI para generar una respuesta en 
# streaming utilizando el modelo "gpt-3.5-turbo" y los mensajes del historial. A medida que se recibe la 
# respuesta, se agrega el contenido generado a la respuesta y se actualiza el contenedor con la respuesta 
# en formato Markdown. Se crea un nuevo mensaje con la respuesta del asistente. Si se encuentra una consulta 
# SQL en la respuesta (utilizando una expresi칩n regular), se extrae la consulta, se establece una conexi칩n 
# con Snowflake, se ejecuta la consulta y se guardan los resultados en el mensaje. Finalmente, se agrega 
# el mensaje del asistente al historial de mensajes.

# Si el 칰ltimo mensaje no es del asistente, se necesita generar una nueva respuesta.
if st.session_state.messages[-1]["role"] != "assistant":  # Verifica si el 칰ltimo mensaje no es del asistente.
    with st.chat_message("assistant"):  # Crea un contenedor para el mensaje del asistente.
        response = ""  # Inicializa la respuesta como una cadena vac칤a.
        resp_container = st.empty()  # Crea un contenedor vac칤o para actualizar din치micamente la respuesta.
        for delta in client.chat.completions.create(  # Llama a la API de OpenAI para generar una respuesta en streaming.
            model="gpt-3.5-turbo",  # Especifica el modelo de lenguaje a utilizar.
            messages=[{"role": m["role"], "content": m["content"]} for m in st.session_state.messages],  # Prepara los mensajes para la API.
            stream=True,  # Habilita el modo de streaming.
        ):
            response += (delta.choices[0].delta.content or "")  # Agrega el contenido generado a la respuesta.
            resp_container.markdown(response)  # Actualiza el contenedor con la respuesta en Markdown.

        message = {"role": "assistant", "content": response}  # Crea un nuevo mensaje con la respuesta del asistente.
        # Analiza la respuesta en busca de una consulta SQL y la ejecuta si est치 disponible.
        sql_match = re.search(r"```sql\n(.*)\n```", response, re.DOTALL)  # Busca una consulta SQL en la respuesta usando una expresi칩n regular.
        if sql_match:  # Si se encuentra una consulta SQL.
            sql = sql_match.group(1)  # Extrae la consulta SQL.
            conn = st.connection("snowflake")  # Establece una conexi칩n con Snowflake.
            message["results"] = conn.query(sql)  # Ejecuta la consulta SQL y guarda los resultados.
            st.dataframe(message["results"])  # Muestra los resultados en un DataFrame.
        st.session_state.messages.append(message)  # Agrega el mensaje del asistente al historial de mensajes.

        # Agrega el bot칩n de exportaci칩n despu칠s de la respuesta
        if "results" in message:
            export_button = st.download_button(
                label="Descargar archivo Word",
                data=to_word(message["results"], prompt),
                file_name="resultados.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
