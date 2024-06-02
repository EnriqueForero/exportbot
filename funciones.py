from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement, parse_xml
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Tomado de la programación de María Paula Diaz para el reporte de la Vice de exportaciones. 

# Definir nombres de estilo como constantes
TITLE_STYLE_NAME = 'MyTitleStyle'
HEADING1_STYLE_NAME = 'MyHeading1'
HEADING2_STYLE_NAME = 'MyHeading2'
TABLE_STYLE_NAME = 'MyTableStyle'
NORMAL_STYLE_NAME = 'Normal'

def estilos(doc):
    """
    Define y aplica estilos personalizados al documento proporcionado.
    
    Args:
    doc (Document): El documento al que se añadirán los estilos.
    """
    # Función auxiliar para verificar y crear estilos si no existen
    def add_style_if_not_exist(name, style_type):
        if name not in doc.styles:
            return doc.styles.add_style(name, style_type)
        return doc.styles[name]

    # Estilo para Título
    title_style = add_style_if_not_exist(TITLE_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(16)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Estilo para Encabezado 1
    heading1_style = add_style_if_not_exist(HEADING1_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
    heading1_font = heading1_style.font
    heading1_font.name = 'Arial'
    heading1_font.size = Pt(14)
    heading1_font.bold = True
    heading1_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Estilo para Encabezado 2
    heading2_style = add_style_if_not_exist(HEADING2_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
    heading2_font = heading2_style.font
    heading2_font.name = 'Arial'
    heading2_font.size = Pt(12)
    heading2_font.bold = True
    heading2_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Estilo para Normal
    normal_style = add_style_if_not_exist(NORMAL_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Estilo para Tabla
    table_style = add_style_if_not_exist(TABLE_STYLE_NAME, WD_STYLE_TYPE.TABLE)
    table_style.font.name = 'Arial'
    table_style.font.size = Pt(10)

def formato_tablas(doc, table):
    """
    Formatea la tabla proporcionada aplicando estilos, colores, alineaciones y bordes.
    
    Args:
    doc (Document): El documento de Word.
    table (Table): La tabla a formatear.
    """
    # Formatear el encabezado de la tabla
    for row in table.rows[:1]:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            shading_elm = parse_xml(r'<w:shd {} w:fill="002060"/>'.format(nsdecls('w')))
            cell._tc.get_or_add_tcPr().append(shading_elm)

    # Formatear el contenido de la tabla
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)

    # Configurar la alineación vertical y horizontal del texto en el centro
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), "center")
            tcPr.append(tcVAlign)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Añadir bordes a la tabla
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), 'auto')
            tcBorders.append(top)
            
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '4')
            left.set(qn('w:space'), '0')
            left.set(qn('w:color'), 'auto')
            tcBorders.append(left)
            
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'auto')
            tcBorders.append(bottom)
            
            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'single')
            right.set(qn('w:sz'), '4')
            right.set(qn('w:space'), '0')
            right.set(qn('w:color'), 'auto')
            tcBorders.append(right)

            tcPr.append(tcBorders)

    # Ajustar el ancho de las columnas automáticamente
    table.autofit = True

    # Establecer el estilo de la tabla para que se ajuste al contenido
    table.style = 'Table Grid'

    # Ajustar la altura de las filas automáticamente según su contenido
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:hRule'), "auto")
        trPr.append(trHeight)

def formatear_numero(valor):
    """
    Formatea un número o una cadena que representa un número.
    
    Args:
    valor (int, float, str): El número o cadena a formatear.
    
    Returns:
    str: El número formateado como cadena.
    """
    try:
        if isinstance(valor, str) and not valor.replace('.', '', 1).isdigit():
            return valor
        valor = float(valor)
        if isinstance(valor, float) and not valor.is_integer():
            return "{:,.2f}".format(valor)
        else:
            return "{:,}".format(int(valor))
    except (ValueError, TypeError):
        return str(valor)

def nsdecls(*prefixes):
    """
    Genera declaraciones de espacio de nombres para su uso en elementos XML.
    
    Args:
    *prefixes (str): Prefijos de espacio de nombres.
    
    Returns:
    str: Declaraciones de espacio de nombres como cadena.
    """
    return ' '.join(['xmlns:{}="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'.format(p) for p in prefixes])